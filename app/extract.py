"""Extract plain text from uploaded CVs so the agent can screen without opening files."""
import io
import logging

log = logging.getLogger(__name__)

MAX_CHARS = 60_000


def extract_text(data: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "pdf":
            text = _pdf(data)
        elif ext == "docx":
            text = _docx(data)
        elif ext == "doc":
            text = _doc_fallback(data)
        else:
            text = ""
    except Exception as exc:  # never block an application on extraction
        log.warning("CV text extraction failed for %s: %s", filename, exc)
        text = ""
    text = "\n".join(line.rstrip() for line in text.splitlines())
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text.strip()[:MAX_CHARS]


def _pdf(data):
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx(data):
    import docx

    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _doc_fallback(data):
    # Legacy .doc: best-effort ASCII salvage (recommend candidates upload PDF/DOCX)
    raw = data.decode("latin-1", errors="ignore")
    printable = "".join(ch if ch.isprintable() or ch in "\n\t" else " " for ch in raw)
    return " ".join(printable.split())
